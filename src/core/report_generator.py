import logging
import markdown
from pathlib import Path
from typing import Dict, Any, Optional

try:
    from docx import Document

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)


class ReportGenerator:
    """
    Generates reports in multiple formats (Markdown, HTML, Docx).
    """

    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def save_report(
        self, content: str, filename: str, formats: list[str] = ["md"]
    ) -> Dict[str, str]:
        """
        Save report in specified formats.
        Returns a dictionary of format -> file_path.
        """
        saved_files = {}
        base_name = filename.rsplit(".", 1)[0]

        for fmt in formats:
            try:
                if fmt == "md":
                    path = self._save_markdown(content, base_name)
                    saved_files["md"] = str(path)
                elif fmt == "html":
                    path = self._save_html(content, base_name)
                    saved_files["html"] = str(path)
                elif fmt == "docx":
                    path = self._save_docx(content, base_name)
                    saved_files["docx"] = str(path)
                else:
                    logger.warning(f"Unsupported format: {fmt}")
            except Exception as e:
                logger.error(f"Failed to save {fmt} report: {e}")

        return saved_files

    def _save_markdown(self, content: str, base_name: str) -> Path:
        path = self.output_dir / f"{base_name}.md"
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        logger.info(f"Saved Markdown report to {path}")
        return path

    def _save_html(self, content: str, base_name: str) -> Path:
        path = self.output_dir / f"{base_name}.html"
        html_content = markdown.markdown(content)
        # Add basic styling
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{base_name}</title>
            <style>
                body {{ font-family: sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
                h1, h2, h3 {{ color: #333; }}
                code {{ background: #f4f4f4; padding: 2px 5px; border-radius: 3px; }}
                pre {{ background: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }}
                a {{ color: #0066cc; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        with open(path, "w", encoding="utf-8") as f:
            f.write(full_html)
        logger.info(f"Saved HTML report to {path}")
        return path

    def _save_docx(self, content: str, base_name: str) -> Path:
        if not HAS_DOCX:
            logger.warning("python-docx not installed, skipping docx generation")
            return Path("")

        path = self.output_dir / f"{base_name}.docx"
        doc = Document()

        # Simple Markdown to Docx conversion (very basic)
        # For a robust solution, we'd need a proper parser.
        # Here we just dump the text, maybe handling headers.

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue

            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

        doc.save(path)
        logger.info(f"Saved Docx report to {path}")
        return path
