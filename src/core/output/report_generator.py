import logging
import markdown
from pathlib import Path
from typing import Dict

try:
    from docx import Document

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from weasyprint import HTML, CSS

    HAS_WEASYPRINT = True
except (ImportError, OSError):
    # OSError occurs when WeasyPrint's native GTK libraries are missing on Windows
    HAS_WEASYPRINT = False
    HTML = None
    CSS = None

# Fallback PDF generator (pure Python, no native deps)
try:
    from fpdf import FPDF

    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False
    FPDF = None

logger = logging.getLogger(__name__)


class ReportGenerator:
    """
    Generates reports in multiple formats (Markdown, HTML, Docx, PDF).
    """

    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def save_report(
        self, content: str, filename: str, formats: list[str] = ["md"]
    ) -> Dict[str, str]:
        """
        Save report in specified formats.
        Returns a dictionary of format -> file_path.
        """
        saved_files = {}
        base_name = filename.rsplit(".", 1)[0]

        for fmt in formats:
            try:
                if fmt == "md":
                    path = self._save_markdown(content, base_name)
                    saved_files["md"] = str(path)
                elif fmt == "html":
                    path = self._save_html(content, base_name)
                    saved_files["html"] = str(path)
                elif fmt == "docx":
                    path = self._save_docx(content, base_name)
                    saved_files["docx"] = str(path)
                elif fmt == "pdf":
                    path = self._save_pdf(content, base_name)
                    if path:
                        saved_files["pdf"] = str(path)
                else:
                    logger.warning(f"Unsupported format: {fmt}")
            except Exception as e:
                logger.error(f"Failed to save {fmt} report: {e}")

        return saved_files

    def _save_markdown(self, content: str, base_name: str) -> Path:
        path = self.output_dir / f"{base_name}.md"
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        logger.info(f"Saved Markdown report to {path}")
        return path

    def _save_html(self, content: str, base_name: str) -> Path:
        path = self.output_dir / f"{base_name}.html"
        html_content = markdown.markdown(content)
        # Add basic styling
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{base_name}</title>
            <style>
                body {{ font-family: sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
                h1, h2, h3 {{ color: #333; }}
                code {{ background: #f4f4f4; padding: 2px 5px; border-radius: 3px; }}
                pre {{ background: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }}
                a {{ color: #0066cc; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        with open(path, "w", encoding="utf-8") as f:
            f.write(full_html)
        logger.info(f"Saved HTML report to {path}")
        return path

    def _save_docx(self, content: str, base_name: str) -> Path:
        if not HAS_DOCX:
            logger.warning("python-docx not installed, skipping docx generation")
            return Path("")

        path = self.output_dir / f"{base_name}.docx"
        doc = Document()

        # Simple Markdown to Docx conversion (very basic)
        # For a robust solution, we'd need a proper parser.
        # Here we just dump the text, maybe handling headers.

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue

            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

        doc.save(path)
        logger.info(f"Saved Docx report to {path}")
        return path

    def _save_pdf(self, content: str, base_name: str) -> Path:
        """
        Convert Markdown content to a professionally styled PDF.

        Uses WeasyPrint for high-quality PDF generation with CSS styling.
        Falls back to fpdf2 (pure Python) if WeasyPrint is not available.
        """
        if not HAS_WEASYPRINT:
            if HAS_FPDF:
                return self._save_pdf_fpdf(content, base_name)
            logger.warning(
                "No PDF library available. Install with: pip install fpdf2"
            )
            return Path("")

        path = self.output_dir / f"{base_name}.pdf"

        # Convert Markdown to HTML with extensions for tables and code
        html_content = markdown.markdown(
            content,
            extensions=["tables", "fenced_code", "toc"],
        )

        # Professional PDF styling
        pdf_css = CSS(
            string="""
            @page {
                size: A4;
                margin: 2cm;
                @top-center {
                    content: string(title);
                    font-size: 10pt;
                    color: #666;
                }
                @bottom-center {
                    content: "Page " counter(page) " of " counter(pages);
                    font-size: 9pt;
                    color: #666;
                }
            }
            body {
                font-family: 'Helvetica Neue', Arial, sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #333;
            }
            h1 {
                string-set: title content();
                color: #1a1a2e;
                border-bottom: 2px solid #4a90d9;
                padding-bottom: 0.3em;
                margin-top: 1.5em;
            }
            h2 {
                color: #16213e;
                border-bottom: 1px solid #ddd;
                padding-bottom: 0.2em;
                margin-top: 1.2em;
            }
            h3 {
                color: #0f3460;
                margin-top: 1em;
            }
            table {
                width: 100%;
                border-collapse: collapse;
                margin: 1em 0;
                font-size: 10pt;
            }
            th, td {
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }
            th {
                background-color: #4a90d9;
                color: white;
            }
            tr:nth-child(even) {
                background-color: #f9f9f9;
            }
            code {
                background: #f4f4f4;
                padding: 2px 5px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
                font-size: 10pt;
            }
            pre {
                background: #f8f8f8;
                border: 1px solid #ddd;
                border-radius: 5px;
                padding: 1em;
                overflow-x: auto;
                font-size: 9pt;
            }
            blockquote {
                border-left: 4px solid #4a90d9;
                margin: 1em 0;
                padding-left: 1em;
                color: #666;
                font-style: italic;
            }
            ul, ol {
                margin: 0.5em 0;
                padding-left: 2em;
            }
            li {
                margin: 0.3em 0;
            }
            a {
                color: #4a90d9;
                text-decoration: none;
            }
            .toc {
                background: #f9f9f9;
                border: 1px solid #ddd;
                padding: 1em;
                margin: 1em 0;
            }
        """
        )

        # Wrap in full HTML document
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{base_name}</title>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        try:
            html_doc = HTML(string=full_html)
            html_doc.write_pdf(path, stylesheets=[pdf_css])
            logger.info(f"Saved PDF report to {path}")
            return path
        except Exception as e:
            logger.error(f"Failed to generate PDF: {e}")
            return Path("")

    def _save_pdf_fpdf(self, content: str, base_name: str) -> Path:
        """
        Fallback PDF generation using fpdf2 (pure Python, no native deps).
        """
        path = self.output_dir / f"{base_name}.pdf"

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)

        for line in content.split("\n"):
            line = line.rstrip()

            # Handle headers
            if line.startswith("# "):
                pdf.set_font("Helvetica", "B", 16)
                pdf.cell(0, 10, line[2:], new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("Helvetica", size=11)
            elif line.startswith("## "):
                pdf.set_font("Helvetica", "B", 14)
                pdf.cell(0, 8, line[3:], new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("Helvetica", size=11)
            elif line.startswith("### "):
                pdf.set_font("Helvetica", "B", 12)
                pdf.cell(0, 7, line[4:], new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("Helvetica", size=11)
            elif line.startswith("- "):
                pdf.cell(0, 6, f"  - {line[2:]}", new_x="LMARGIN", new_y="NEXT")
            elif line.strip() == "":
                pdf.ln(3)
            else:
                # Multi-cell for wrapping long lines
                pdf.multi_cell(0, 6, line)

        pdf.output(path)
        logger.info(f"Saved PDF report to {path} (using fpdf2)")
        return path
